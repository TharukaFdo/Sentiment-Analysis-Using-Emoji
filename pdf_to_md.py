"""
PDF/DOCX to Markdown Converter
Converts PDF or DOCX files to Markdown format.

Usage:
    python pdf_to_md.py <input.pdf|input.docx> [output.md]
    python pdf_to_md.py <input_dir> [output_dir] [--recursive]
    
If output path is not specified, creates a .md file with the same name as the input.
If a directory is provided, converts all supported files in that directory into
Markdown and writes them to a `markdown_exports` folder by default.

Requirements:
    pip install pymupdf python-docx
"""

import argparse
import sys
import os
import re
from urllib.parse import quote

# PDF support
import fitz  # PyMuPDF

# DOCX support
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.table import Table as DocxTable


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


def normalize_text_chars(text):
    """Normalize common DOCX text artifacts for cleaner Markdown."""
    if text is None:
        return ""
    return (
        str(text)
        .replace("\xa0", " ")   # non-breaking space
        .replace("\u200b", "")  # zero-width space
        .replace("\ufeff", "")  # BOM
    )


def escape_accidental_markdown_headings(text):
    """
    Escape lines that accidentally start with Markdown heading syntax.
    This prevents source content like '# comment' from becoming a heading.
    """
    escaped_lines = []
    for line in text.split("\n"):
        if re.match(r'^\s*#{1,6}\s+', line):
            first_hash = line.find('#')
            line = line[:first_hash] + "\\" + line[first_hash:]
        escaped_lines.append(line)
    return "\n".join(escaped_lines)


def sanitize_markdown_body_text(text):
    """Prepare paragraph/list text for Markdown body output."""
    cleaned = normalize_text_chars(text)
    cleaned_lines = []
    for line in cleaned.split("\n"):
        # Remove stray trailing quotes after executable/library file names.
        line = re.sub(r"(\.(?:exe|dll|so|dylib))['\"]+\s*$", r"\1", line, flags=re.IGNORECASE)
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines)
    return escape_accidental_markdown_headings(cleaned)


def normalize_for_comparison(text):
    """Normalize text for robust equality checks."""
    return re.sub(r'\s+', ' ', normalize_text_chars(text)).strip().lower()


def to_markdown_path(*parts):
    """Build Markdown-friendly relative paths with forward slashes."""
    encoded_parts = []
    for part in parts:
        if part is None:
            continue
        for segment in str(part).replace("\\", "/").split("/"):
            if not segment:
                continue
            encoded_parts.append(quote(segment, safe="._-()"))
    return "/".join(encoded_parts)


def slugify_markdown_anchor(text):
    """Create a GitHub-style-ish anchor slug for TOC links."""
    slug = normalize_text_chars(text).strip().lower()
    slug = re.sub(r'[`*_~]', '', slug)
    slug = re.sub(r'[^\w\s-]', '', slug, flags=re.UNICODE)
    slug = re.sub(r'\s+', '-', slug)
    slug = re.sub(r'-+', '-', slug).strip('-')
    return slug


def extract_images(page, doc, output_dir, page_num):
    """Extract images from a PDF page and save them."""
    images = []
    image_list = page.get_images(full=True)
    
    for img_index, img_info in enumerate(image_list):
        xref = img_info[0]
        try:
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            # Skip very small images — likely decorative icons or artefacts.
            # A threshold of 2 KB filters out bullets, checkmarks, etc.
            if len(image_bytes) < 2048:
                continue

            # Create image filename
            image_filename = f"page{page_num + 1}_img{img_index + 1}.{image_ext}"
            image_path = os.path.join(output_dir, image_filename)
            
            # Save image
            with open(image_path, "wb") as img_file:
                img_file.write(image_bytes)
            
            images.append(image_filename)
        except Exception as e:
            print(f"Warning: Could not extract image {img_index} from page {page_num + 1}: {e}")
    
    return images


def is_page_header_or_footer(block_text, page_height, block_y, doc_title=""):
    """
    Heuristic: detect running headers/footers that repeat on every page.
    Blocks in the top ~8% or bottom ~8% of the page that look like
    journal metadata are treated as headers/footers and suppressed.
    """
    margin = page_height * 0.08
    if block_y > margin and block_y < (page_height - margin):
        return False

    # Common journal/conference header/footer patterns
    patterns = [
        r'ACM Trans\.',
        r'Vol\.\s*\d+',
        r'Article\s*\d+',
        r'Publication date',
        r'^\d+:\d+\s*[•·]',          # e.g. "102:2 •"
        r'^\s*\d+\s*$',               # lone page number
    ]
    for pat in patterns:
        if re.search(pat, block_text, re.IGNORECASE):
            return True
    return False


def rejoin_hyphenated_lines(text):
    """
    Rejoin words that were hyphenated across column/line breaks by the PDF
    layout engine.  Only merges when the hyphen falls at end-of-line and
    the next non-empty line starts with a lower-case letter (or continues a
    word), which avoids destroying intentional hyphens in compound words
    that sit in the middle of a sentence.
    """
    # Pattern: word- \n lowercase-continuation
    text = re.sub(r'(\w)-\n([a-z])', r'\1\2', text)
    return text


def clean_text(text):
    """Clean and format extracted text."""
    # Rejoin PDF column-break hyphenations before other processing
    text = rejoin_hyphenated_lines(text)

    # Remove excessive whitespace while preserving paragraph breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove trailing whitespace from lines
    lines = [line.rstrip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text


# ---------------------------------------------------------------------------
# Heading detection — PDF version
# ---------------------------------------------------------------------------

# Short tokens that look like acronyms / labels but are NOT document headings.
_HEADING_BLACKLIST = {
    'SDF', 'NRC', 'NeRF', 'MLP', 'GPU', 'CPU', 'RGB', 'HDR',
    'TOC', 'URL', 'DOI', 'PDF', 'API', 'ACM', 'IEEE',
}

# Patterns that indicate the line is running metadata, not a heading.
_METADATA_PATTERNS = [
    r'^\d{4}-\d{4}/',          # ISSN-like strings, e.g. 0730-0301/2022/…
    r'https?://',               # URLs
    r'doi\.org',
    r'@\w+\.',                  # email fragments
    r'©',                       # copyright
    r'\$\d',                    # price / licence
    r'ACM Trans\.',
    r'Vol\.\s*\d+',
    r'Publication date',
    r'^\d+:\d+',               # page:article numbers like "102:3"
]


def _is_metadata_line(line):
    for pat in _METADATA_PATTERNS:
        if re.search(pat, line, re.IGNORECASE):
            return True
    return False


def detect_headings(text):
    """
    Attempt to detect and format headings based on text patterns.

    Key improvements over the original:
    * Blacklisted short tokens (acronyms, labels) are never promoted.
    * Metadata/footer lines are never promoted.
    * Numbered appendix sections like "E.2 Occupancy Grids" are correctly
      detected as headings without the sub-label alone becoming a heading.
    * The all-caps heuristic requires at least two words to reduce noise.
    """
    lines = text.split('\n')
    result = []
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        if not stripped:
            result.append(line)
            continue

        # Never turn metadata lines into headings.
        if _is_metadata_line(stripped):
            result.append(line)
            continue

        is_heading = False
        heading_level = 2

        # --- Numbered sections: "1.", "1.1", "A.2", "E.3 Title" etc. ---
        numbered = re.match(
            r'^([A-Z]|\d+)(\.\d+)*\.?\s+[A-Z]\w', stripped
        )
        if numbered and len(stripped) < 120:
            is_heading = True
            heading_level = 2

        # --- Chapter/Section/Part keywords ---
        elif re.match(r'^(Chapter|Section|Part)\s+\d+', stripped, re.I):
            is_heading = True
            heading_level = 2

        # --- All-caps lines: require ≥2 words to reduce acronym false positives ---
        elif (
            stripped.isupper()
            and len(stripped) < 100
            and len(stripped.split()) >= 2       # ← was <=10, now requires ≥2 words
            and stripped not in _HEADING_BLACKLIST
            and not _is_metadata_line(stripped)
        ):
            is_heading = True
            heading_level = 1

        # --- Short lines without terminal punctuation followed by a blank line ---
        elif (
            len(stripped) < 80
            and not stripped.endswith(('.', ',', ';', ':', '!', '?'))
            and len(stripped.split()) >= 2       # at least two words
            and len(stripped.split()) <= 8
            and i + 1 < len(lines)
            and not lines[i + 1].strip()
            and stripped[0].isupper()
            and stripped not in _HEADING_BLACKLIST
            and not _is_metadata_line(stripped)
            # Must contain at least one "normal" word (not all symbols/digits)
            and re.search(r'[A-Za-z]{3,}', stripped)
        ):
            is_heading = True
            heading_level = 2

        if is_heading:
            result.append(f"{'#' * heading_level} {stripped}")
        else:
            result.append(line)
    
    return '\n'.join(result)


def extract_tables(page):
    """Attempt to extract tables from a page."""
    tables = []
    try:
        # PyMuPDF can find tables in newer versions
        if hasattr(page, 'find_tables'):
            found_tables = page.find_tables()
            for table in found_tables:
                table_data = table.extract()
                if table_data:
                    tables.append((table.bbox, table_data))
    except Exception:
        pass  # Table extraction not available or failed
    return tables


def format_table_as_markdown(table_data):
    """Convert table data to Markdown format."""
    if not table_data or not table_data[0]:
        return ""
    
    # Flatten multi-line cell content and clean None values
    def clean_cell(cell):
        if cell is None:
            return ""
        return re.sub(r'\s+', ' ', str(cell)).strip()

    md_lines = []
    
    # Header row
    header = [clean_cell(c) for c in table_data[0]]
    if not any(header):          # skip completely empty header rows
        if len(table_data) < 2:
            return ""
        header = [clean_cell(c) for c in table_data[1]]
        table_data = table_data[2:]
    else:
        table_data = table_data[1:]

    md_lines.append("| " + " | ".join(header) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    
    # Data rows
    for row in table_data:
        cells = [clean_cell(c) for c in row]
        # Ensure row has same number of cells as header
        while len(cells) < len(header):
            cells.append("")
        md_lines.append("| " + " | ".join(cells[:len(header)]) + " |")
    
    return "\n".join(md_lines)


# =============================================================================
# DOCX Processing Functions
# =============================================================================

def extract_docx_images(doc, output_dir):
    """Extract images from a DOCX document."""
    images = []
    rel_id_to_filename = {}
    image_counter = 0
    
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                image_counter += 1
                image_data = rel.target_part.blob
                
                # Get image extension from content type
                content_type = rel.target_part.content_type
                ext_map = {
                    'image/png': 'png',
                    'image/jpeg': 'jpg',
                    'image/gif': 'gif',
                    'image/bmp': 'bmp',
                    'image/tiff': 'tiff',
                    'image/webp': 'webp'
                }
                ext = ext_map.get(content_type, 'png')
                
                image_filename = f"image_{image_counter}.{ext}"
                image_path = os.path.join(output_dir, image_filename)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_data)
                
                images.append(image_filename)
                rel_id_to_filename[rel_id] = image_filename
            except Exception as e:
                print(f"Warning: Could not extract image: {e}")
    
    return images, rel_id_to_filename


def docx_table_to_markdown(table):
    """Convert a DOCX table to Markdown format."""
    rows = []
    for row in table.rows:
        cells = [normalize_text_chars(cell.text).strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return ""
    
    return format_table_as_markdown(rows)


def get_paragraph_style_level(paragraph):
    """Determine heading level from paragraph style."""
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    
    # Check for heading styles
    if 'heading 1' in style_name or style_name == 'title':
        return 1
    elif 'heading 2' in style_name:
        return 2
    elif 'heading 3' in style_name:
        return 3
    elif 'heading 4' in style_name:
        return 4
    elif 'heading 5' in style_name:
        return 5
    elif 'heading 6' in style_name:
        return 6
    
    return 0  # Not a heading


def format_paragraph_text(paragraph):
    """Format paragraph text with inline styles (bold, italic)."""
    text_parts = []
    
    for run in paragraph.runs:
        text = normalize_text_chars(run.text)
        if not text:
            continue
        
        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        text_parts.append(text)

    formatted_text = "".join(text_parts)

    # Fallback for field-generated paragraphs (e.g., Word List of Figures/Tables),
    # where runs can be empty while paragraph.text still contains visible content.
    if not formatted_text.strip():
        return normalize_text_chars(paragraph.text)

    return formatted_text


def paragraph_has_numbering(paragraph):
    """Check whether a paragraph has Word numbering metadata."""
    try:
        ppr = paragraph._element.pPr
        return ppr is not None and ppr.numPr is not None
    except AttributeError:
        return False


def get_xml_paragraph_style(p_el):
    """Get paragraph style name from a raw XML paragraph element."""
    p_pr = p_el.find(qn('w:pPr'))
    if p_pr is None:
        return ""

    p_style = p_pr.find(qn('w:pStyle'))
    if p_style is None:
        return ""

    return p_style.get(qn('w:val'), "")


def xml_text_with_tabs(el):
    """Extract text from raw XML, preserving tab separators."""
    parts = []
    for node in el.iter():
        tag = node.tag.split('}')[-1] if isinstance(node.tag, str) else ""
        if tag == 't':
            parts.append(node.text or "")
        elif tag == 'tab':
            parts.append('\t')
        elif tag == 'br':
            parts.append('\n')
    return "".join(parts)


def extract_toc_from_sdt(sdt_el):
    """Extract Table of Contents entries from a DOCX content control (w:sdt)."""
    paragraphs = [el for el in sdt_el.iter() if isinstance(el.tag, str) and el.tag.endswith('}p')]
    if not paragraphs:
        return []

    # Only handle TOC content controls.
    has_toc_styles = any(get_xml_paragraph_style(p).lower().startswith('toc') for p in paragraphs)
    if not has_toc_styles:
        return []

    toc_lines = []
    toc_heading_emitted = False

    for p_el in paragraphs:
        style = get_xml_paragraph_style(p_el)
        if not style:
            continue

        style_lower = style.lower()
        line = normalize_text_chars(xml_text_with_tabs(p_el)).strip()
        if not line:
            continue

        if style_lower == 'tocheading':
            toc_lines.append(f"# {line}")
            toc_lines.append("")
            toc_heading_emitted = True
            continue

        if not style_lower.startswith('toc'):
            continue

        parts = [part.strip() for part in line.split('\t') if part.strip()]
        if not parts:
            continue

        label = ""
        page = ""
        if len(parts) >= 3:
            number = parts[0]
            title = parts[1]
            page = parts[-1]
            label = f"{number} {title}".strip()
        elif len(parts) == 2:
            label = parts[0]
            page = parts[1]
        else:
            label = parts[0]

        anchor = slugify_markdown_anchor(label)
        if anchor:
            entry = f"- [{label}](#{anchor})"
        else:
            entry = f"- {label}"
        if page:
            entry += f" (p. {page})"

        level_match = re.match(r'toc(\d+)', style_lower)
        level = int(level_match.group(1)) if level_match else 1
        indent = "  " * max(level - 1, 0)
        toc_lines.append(f"{indent}{entry}")

    if toc_lines and not toc_heading_emitted:
        toc_lines.insert(0, "# Table of Contents")
        toc_lines.insert(1, "")

    if toc_lines:
        toc_lines.append("")

    return toc_lines


def is_list_paragraph(paragraph):
    """Check if paragraph is a list item."""
    # Check for numbering
    if paragraph_has_numbering(paragraph):
        return True
    
    # Check style name
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if 'list' in style_name or 'bullet' in style_name:
        return True
    
    return False


def get_paragraph_image_rel_ids(paragraph):
    """Get embedded image relationship IDs from a paragraph, in order."""
    rel_ids = []
    for run in paragraph.runs:
        try:
            embedded = run._element.xpath('.//*[local-name()="blip"]/@*[local-name()="embed"]')
        except Exception:
            embedded = []
        for rel_id in embedded:
            if rel_id not in rel_ids:
                rel_ids.append(rel_id)
    return rel_ids


def docx_to_markdown(docx_path, output_path=None, extract_images_flag=True):
    """
    Convert a DOCX file to Markdown.
    
    Args:
        docx_path: Path to the input DOCX file
        output_path: Path for the output Markdown file (optional)
        extract_images_flag: Whether to extract and include images
    
    Returns:
        Path to the created Markdown file
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    # Determine output path
    if output_path is None:
        base_name = os.path.splitext(docx_path)[0]
        output_path = base_name + ".md"
    
    # Create images directory if needed
    images_dir = None
    if extract_images_flag:
        images_dir = os.path.splitext(output_path)[0] + "_images"
    
    # Open DOCX
    doc = Document(docx_path)
    
    markdown_content = []
    
    # Try to get title from core properties
    try:
        title = doc.core_properties.title
    except:
        title = None
    
    if not title:
        title = os.path.splitext(os.path.basename(docx_path))[0]
    
    title = normalize_text_chars(title).strip()
    title_compare = normalize_for_comparison(title)
    markdown_content.append(f"# {title}\n")
    
    print(f"Processing DOCX document...")
    
    # Extract images first
    extracted_images = []
    rel_id_to_filename = {}
    if extract_images_flag:
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)
        extracted_images, rel_id_to_filename = extract_docx_images(doc, images_dir)
    
    # Process document body
    list_counter = 0
    prev_was_list = False
    heading_counters = [0, 0, 0, 0, 0, 0]
    numbered_heading_mode = False
    pending_title_dedup = True
    used_docx_images = set()
    
    for element in doc.element.body:
        # Handle tables
        if element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element is element:
                    table_md = docx_table_to_markdown(table)
                    if table_md:
                        markdown_content.append(f"\n{table_md}\n")
                    break
        
        # Handle structured document tags (e.g., Table of Contents block)
        elif element.tag.endswith('sdt'):
            toc_lines = extract_toc_from_sdt(element)
            if toc_lines:
                markdown_content.extend(toc_lines)
        
        # Handle paragraphs
        elif element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element is element:
                    text = format_paragraph_text(para)
                    paragraph_image_rel_ids = get_paragraph_image_rel_ids(para) if extract_images_flag else []
                    text = normalize_text_chars(text)
                    
                    if not text.strip():
                        if prev_was_list:
                            list_counter = 0
                            prev_was_list = False
                        if paragraph_image_rel_ids:
                            for rel_id in paragraph_image_rel_ids:
                                img_name = rel_id_to_filename.get(rel_id)
                                if not img_name:
                                    continue
                                rel_path = to_markdown_path(os.path.basename(images_dir), img_name)
                                markdown_content.append(f"![{img_name}]({rel_path})")
                                used_docx_images.add(img_name)
                        else:
                            markdown_content.append("")
                        break

                    # Avoid duplicate cover title when the first body paragraph repeats the document title.
                    if pending_title_dedup:
                        pending_title_dedup = False
                        if normalize_for_comparison(text) == title_compare:
                            break
                    
                    # Check for heading
                    heading_level = get_paragraph_style_level(para)
                    
                    if heading_level > 0:
                        heading_text = normalize_text_chars(text).strip()

                        if paragraph_has_numbering(para):
                            numbered_heading_mode = True

                        if numbered_heading_mode:
                            heading_counters[heading_level - 1] += 1
                            for idx in range(heading_level, len(heading_counters)):
                                heading_counters[idx] = 0

                            section_number = ".".join(
                                str(n) for n in heading_counters[:heading_level] if n > 0
                            )
                            markdown_content.append(
                                f"\n{'#' * heading_level} {section_number} {heading_text}\n"
                            )
                        else:
                            markdown_content.append(f"\n{'#' * heading_level} {heading_text}\n")
                        prev_was_list = False
                        list_counter = 0
                    # Check for list
                    elif is_list_paragraph(para):
                        style_name = para.style.name.lower() if para.style else ""
                        item_text = sanitize_markdown_body_text(text).strip()
                        if 'number' in style_name or 'ordered' in style_name:
                            list_counter += 1
                            markdown_content.append(f"{list_counter}. {item_text}")
                        else:
                            markdown_content.append(f"- {item_text}")
                        prev_was_list = True
                    else:
                        markdown_content.append(sanitize_markdown_body_text(text))
                        prev_was_list = False
                        list_counter = 0

                    if paragraph_image_rel_ids:
                        for rel_id in paragraph_image_rel_ids:
                            img_name = rel_id_to_filename.get(rel_id)
                            if not img_name:
                                continue
                            rel_path = to_markdown_path(os.path.basename(images_dir), img_name)
                            markdown_content.append(f"![{img_name}]({rel_path})")
                            used_docx_images.add(img_name)
                    break
    
    # Add any images that were not referenced inline.
    if extracted_images:
        remaining_images = [img for img in extracted_images if img not in used_docx_images]
        if remaining_images:
            markdown_content.append("\n## Images\n")
            for img_name in remaining_images:
                rel_path = to_markdown_path(os.path.basename(images_dir), img_name)
                markdown_content.append(f"![{img_name}]({rel_path})")
    
    # Write markdown file
    full_content = "\n".join(markdown_content)
    full_content = re.sub(r'\n{3,}', '\n\n', full_content).strip() + "\n"
    
    with open(output_path, "w", encoding="utf-8", errors="replace") as f:
        f.write(full_content)
    
    print(f"\nConversion complete!")
    print(f"Output: {output_path}")
    if images_dir and os.path.exists(images_dir) and extracted_images:
        print(f"Images: {images_dir}")
    
    return output_path


# =============================================================================
# PDF Processing Functions
# =============================================================================

def _get_page_columns(page):
    """
    Detect whether a page uses a two-column layout and return the x-coordinate
    of the column divider.  Returns None for single-column pages.

    Strategy: look at the horizontal centres of all text blocks.  If there is
    a clear gap in the middle third of the page width, the page is two-column.
    """
    blocks = page.get_text("blocks")   # (x0, y0, x1, y1, text, block_no, block_type)
    if not blocks:
        return None

    page_width = page.rect.width
    third = page_width / 3

    # Collect horizontal mid-points of blocks that contain text
    mids = []
    for b in blocks:
        if b[6] != 0:   # skip image blocks
            continue
        bx0, bx1 = b[0], b[2]
        mid = (bx0 + bx1) / 2
        mids.append(mid)

    if len(mids) < 4:
        return None

    # Count how many block centres fall in the left third, middle third, right third
    left   = sum(1 for m in mids if m < third)
    middle = sum(1 for m in mids if third <= m <= 2 * third)
    right  = sum(1 for m in mids if m > 2 * third)

    # Two-column heuristic: most blocks are NOT in the middle third
    if left >= 2 and right >= 2 and middle <= max(1, len(mids) * 0.15):
        return page_width / 2   # divider roughly at centre
    return None


def _extract_text_two_column(page, divider_x):
    """
    Extract text from a two-column page by reading the left column first,
    then the right column, each sorted top-to-bottom.
    """
    blocks = page.get_text("blocks")
    page_height = page.rect.height

    left_blocks  = []
    right_blocks = []

    for b in blocks:
        if b[6] != 0:
            continue
        bx0, by0, bx1 = b[0], b[1], b[2]
        mid_x = (bx0 + bx1) / 2
        if mid_x < divider_x:
            left_blocks.append((by0, b[4]))
        else:
            right_blocks.append((by0, b[4]))

    left_blocks.sort(key=lambda x: x[0])
    right_blocks.sort(key=lambda x: x[0])

    text_parts = []
    for _, txt in left_blocks:
        text_parts.append(txt.strip())
    for _, txt in right_blocks:
        text_parts.append(txt.strip())

    return "\n\n".join(p for p in text_parts if p)


def _extract_text_single_column(page):
    """Extract text from a single-column page, sorted top-to-bottom."""
    blocks = page.get_text("blocks")
    blocks = [b for b in blocks if b[6] == 0]   # text blocks only
    blocks.sort(key=lambda b: b[1])              # sort by y0
    return "\n\n".join(b[4].strip() for b in blocks if b[4].strip())


def _filter_header_footer_blocks(page, blocks):
    """Remove blocks that are likely running headers or footers."""
    page_height = page.rect.height
    filtered = []
    for b in blocks:
        if b[6] != 0:
            filtered.append(b)
            continue
        # y-midpoint of the block
        by_mid = (b[1] + b[3]) / 2
        if is_page_header_or_footer(b[4], page_height, by_mid):
            continue
        filtered.append(b)
    return filtered


def pdf_to_markdown(pdf_path, output_path=None, extract_images_flag=True):
    """
    Convert a PDF file to Markdown.
    
    Args:
        pdf_path: Path to the input PDF file
        output_path: Path for the output Markdown file (optional)
        extract_images_flag: Whether to extract and include images
    
    Returns:
        Path to the created Markdown file
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    # Determine output path
    if output_path is None:
        base_name = os.path.splitext(pdf_path)[0]
        output_path = base_name + ".md"
    
    # Create images directory if needed
    images_dir = None
    if extract_images_flag:
        images_dir = os.path.splitext(output_path)[0] + "_images"
    
    # Open PDF
    doc = fitz.open(pdf_path)
    
    markdown_content = []
    
    # Add title from PDF metadata or filename
    metadata = doc.metadata
    title = metadata.get('title', '') if metadata else ''
    if not title:
        title = os.path.splitext(os.path.basename(pdf_path))[0]
    
    markdown_content.append(f"# {title}\n")
    
    # Process each page
    total_pages = len(doc)
    print(f"Processing {total_pages} pages...")

    # Pre-extract tables per page so we can skip those regions in text extraction
    # (prevents table content from appearing twice — once as text, once as table)
    page_tables = {}  # page_num -> list of (bbox, markdown_string)
    for page_num in range(total_pages):
        page = doc[page_num]
        raw_tables = extract_tables(page)
        rendered = []
        for bbox, tdata in raw_tables:
            md = format_table_as_markdown(tdata)
            if md:
                rendered.append((bbox, md))
        page_tables[page_num] = rendered

    for page_num in range(total_pages):
        page = doc[page_num]
        page_height = page.rect.height

        # ------------------------------------------------------------------
        # 1. Determine layout (single vs two-column)
        # ------------------------------------------------------------------
        divider_x = _get_page_columns(page)

        # ------------------------------------------------------------------
        # 2. Extract and filter text blocks
        # ------------------------------------------------------------------
        if divider_x is not None:
            raw_text = _extract_text_two_column(page, divider_x)
        else:
            # Use block-level extraction so we can filter headers/footers
            all_blocks = page.get_text("blocks")
            filtered   = _filter_header_footer_blocks(page, all_blocks)
            filtered.sort(key=lambda b: b[1])
            raw_text   = "\n\n".join(b[4].strip() for b in filtered if b[6] == 0 and b[4].strip())

        text = clean_text(raw_text)
        text = detect_headings(text)

        # ------------------------------------------------------------------
        # 3. Build page section (no page-number banners — they add noise)
        # ------------------------------------------------------------------
        if text.strip():
            markdown_content.append(text)
            markdown_content.append("")

        # ------------------------------------------------------------------
        # 4. Insert tables extracted from this page
        # ------------------------------------------------------------------
        for _bbox, table_md in page_tables[page_num]:
            markdown_content.append(f"\n{table_md}\n")

        # ------------------------------------------------------------------
        # 5. Extract and insert images
        # ------------------------------------------------------------------
        if extract_images_flag:
            if not os.path.exists(images_dir):
                os.makedirs(images_dir)
            
            images = extract_images(page, doc, images_dir, page_num)
            for img_name in images:
                rel_path = to_markdown_path(os.path.basename(images_dir), img_name)
                markdown_content.append(f"\n![{img_name}]({rel_path})\n")
        
        print(f"  Processed page {page_num + 1}/{total_pages}")
    
    doc.close()
    
    # Final cleanup
    full_content = "\n".join(markdown_content)
    full_content = re.sub(r'\n{3,}', '\n\n', full_content).strip() + "\n"

    with open(output_path, "w", encoding="utf-8", errors="replace") as f:
        f.write(full_content)
    
    print(f"\nConversion complete!")
    print(f"Output: {output_path}")
    if images_dir and os.path.exists(images_dir):
        print(f"Images: {images_dir}")
    
    return output_path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def iter_supported_files(input_dir, recursive=False):
    """Yield supported file paths from a directory."""
    if recursive:
        walker = os.walk(input_dir)
        for root, _dirs, files in walker:
            for name in files:
                ext = os.path.splitext(name)[1].lower()
                if ext in SUPPORTED_EXTENSIONS:
                    yield os.path.join(root, name)
    else:
        for name in os.listdir(input_dir):
            file_path = os.path.join(input_dir, name)
            if not os.path.isfile(file_path):
                continue
            ext = os.path.splitext(name)[1].lower()
            if ext in SUPPORTED_EXTENSIONS:
                yield file_path


def convert_file(input_path, output_path=None, extract_images_flag=True):
    """Convert a single supported file to Markdown."""
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        return pdf_to_markdown(input_path, output_path, extract_images_flag=extract_images_flag)
    if ext in [".docx", ".doc"]:
        if ext == ".doc":
            print("Warning: .doc files may not be fully supported. Consider converting to .docx first.")
        return docx_to_markdown(input_path, output_path, extract_images_flag=extract_images_flag)
    raise ValueError(f"Unsupported file type '{ext}'. Supported: .pdf, .docx, .doc")


def convert_directory(input_dir, output_dir=None, recursive=False, extract_images_flag=True):
    """Convert all supported files in a directory."""
    if output_dir is None:
        output_dir = os.path.join(input_dir, "markdown_exports")

    source_files = list(iter_supported_files(input_dir, recursive=recursive))
    if not source_files:
        raise FileNotFoundError(
            f"No supported files found in directory: {input_dir}"
        )

    os.makedirs(output_dir, exist_ok=True)

    converted_outputs = []
    for source_path in source_files:
        relative_path = os.path.relpath(source_path, input_dir)
        relative_base, _ext = os.path.splitext(relative_path)
        target_path = os.path.join(output_dir, relative_base + ".md")
        os.makedirs(os.path.dirname(target_path), exist_ok=True)

        print(f"\n=== Converting: {relative_path} ===")
        converted_outputs.append(
            convert_file(
                source_path,
                output_path=target_path,
                extract_images_flag=extract_images_flag,
            )
        )

    print(f"\nBatch conversion complete. Converted {len(converted_outputs)} file(s).")
    print(f"Output directory: {output_dir}")
    return converted_outputs


def main():
    parser = argparse.ArgumentParser(description="Convert PDF or DOCX files to Markdown.")
    parser.add_argument("input_path", help="Input file or directory")
    parser.add_argument(
        "output_path",
        nargs="?",
        help="Output Markdown file for single-file mode, or output directory for batch mode",
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Recursively convert supported files when the input path is a directory",
    )
    parser.add_argument(
        "--no-images",
        action="store_true",
        help="Disable image extraction",
    )

    args = parser.parse_args()

    if not args.input_path:
        print(__doc__)
        print("\nError: Please provide a PDF or DOCX file path.")
        print("Example: python pdf_to_md.py document.pdf")
        print("         python pdf_to_md.py document.docx")
        sys.exit(1)

    input_path = args.input_path
    output_path = args.output_path
    extract_images_flag = not args.no_images

    try:
        if os.path.isdir(input_path):
            convert_directory(
                input_path,
                output_dir=output_path,
                recursive=args.recursive,
                extract_images_flag=extract_images_flag,
            )
        else:
            convert_file(
                input_path,
                output_path=output_path,
                extract_images_flag=extract_images_flag,
            )
    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error converting file: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
